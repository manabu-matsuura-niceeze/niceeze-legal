"""SURPLUS_SHIFT_SRS_v1.0.docx 生成スクリプト"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = "docs/SURPLUS_SHIFT/SURPLUS_SHIFT_SRS_v1.0.docx"
DARK_NAVY = RGBColor(0x1a, 0x3a, 0x5c)
ORANGE    = RGBColor(0xf5, 0xa6, 0x23)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)


def set_margins(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Mm(20)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_table_width(table, width_dxa=9026):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12) if level == 1 else Pt(10.5)
    run.font.color.rgb = DARK_NAVY


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)


def add_full_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_width(table, 9026)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hdr.cells[i], '1a3a5c')
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F0F4F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            set_cell_bg(tr.cells[ci], bg)
            run = tr.cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('SURPLUS_SHIFT 余剰在庫転換システム  |  ソフトウェア要件仕様書 v1.0  |  © 2026 株式会社NiceEze  Confidential')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)


def build():
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc = Document()
    set_margins(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SURPLUS_SHIFT — ソフトウェア要件仕様書 (SRS) v1.0')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('余剰在庫転換システム（SURPLUS_SHIFT）　v1.0').font.size = Pt(12)

    # 1. 文書管理表
    add_heading(doc, '1. 文書管理表')
    add_full_table(doc, ['項目', '内容'], [
        ['文書名', 'SURPLUS_SHIFT ソフトウェア要件仕様書'],
        ['バージョン', 'v1.0'],
        ['作成日', '2026-06-05'],
        ['作成者', 'NiceEze 自律COO'],
        ['承認者', '代表取締役CEO 松浦 学'],
        ['根拠文書', 'SURPLUS_SHIFT_BRD_v1.0.docx / PLAN-20260605-001'],
    ])

    # 2. システム概要
    add_heading(doc, '2. システム概要・アーキテクチャ')
    add_body(doc,
        'SURPLUS_SHIFTはGCPサーバーレス構成を採用した余剰在庫転換プラットフォームである。'
        'Cloud RunによるコンテナベースのAPIサーバー、Firestoreによるリアルタイムデータ管理、'
        'Cloud Functionsによるイベント駆動処理を組み合わせ、'
        'スケーラブルかつコスト最適化されたシステムを実現する。'
        'フロントエンドはVanilla JS + IndexedDB（niceeze_cache_v142）によるオフライン対応PWAとし、'
        'モバイルファーストUIで売り手・買い手双方の操作性を確保する。'
    )
    add_full_table(doc, ['コンポーネント', '役割', '技術'], [
        ['売り手ダッシュボード', '余剰在庫登録・管理・転換状況確認', 'HTML/JS / IndexedDB v142 / Firestore'],
        ['買い手ポータル', 'マッチング候補表示・交渉・購入確定', 'PWA / IndexedDB v142 / Firestore'],
        ['Cloud Run APIサーバー', 'マッチングロジック / 契約書生成 / 認証', 'Python / FastAPI'],
        ['Cloud Functions', 'マッチング非同期処理 / メール通知 / 定期バッチ', 'Python 3.11'],
        ['Firestore', 'リアルタイムDB（在庫・マッチング・取引データ）', 'GCP Firestore ネイティブモード'],
        ['BigQuery', '取引履歴・分析レポート月次アーカイブ', 'bigquery_surplus_pipeline.py'],
        ['Secret Manager', 'APIキー・DB認証情報管理', 'GCP Secret Manager'],
    ])

    # 3. 機能要件
    add_heading(doc, '3. 機能要件一覧（SURPLUS-001〜005詳細仕様）')
    add_full_table(doc,
        ['ID', '機能名', '詳細仕様', '優先度', '画面ID'],
        [
            ['SURPLUS-001', '売り手ダッシュボード',
             '在庫ID自動採番 / 商品カテゴリ選択（食品/電子部品/アパレル/その他）/ '
             '数量・仕入価格・希望転換価格・賞味期限（食品のみ）入力 / ステータス管理（登録中/マッチング中/転換済/廃棄）',
             'Must', 'SURPLUS-UI-001'],
            ['SURPLUS-002', '買い手検索マッチング',
             '買い手プロファイル（希望カテゴリ/地域/予算/最低購入数量）とのスコアリングマッチング / '
             'マッチングスコア算出（カテゴリ一致50点 + 価格適合30点 + 地域20点）/ 上位5件提示',
             'Must', 'SURPLUS-UI-002'],
            ['SURPLUS-003', '価格交渉支援',
             '市場参照価格の自動取得（Cloud Function経由） / '
             '推奨転換価格（仕入価格×0.6〜0.8）の提示 / '
             '交渉チャット履歴記録（Firestore） / 交渉期限管理（デフォルト72時間）',
             'Should', 'SURPLUS-UI-003'],
            ['SURPLUS-004', '転換実行・契約書生成',
             '取引確定後の売買契約書自動生成（python-docx）/ '
             'タイムスタンプ付与・SHA-256ハッシュ署名 / '
             'PDF/Word形式でCloud Storageに保存 / 双方へのメール送付',
             'Must', 'SURPLUS-UI-004'],
            ['SURPLUS-005', '履歴・分析レポート',
             '転換実績（件数/金額/転換率）のダッシュボード表示 / '
             'BigQueryによる月次集計 / CSV/Excel形式エクスポート / '
             'カテゴリ別・期間別フィルタリング',
             'Should', 'SURPLUS-UI-005'],
        ]
    )

    # 4. 非機能要件
    add_heading(doc, '4. 非機能要件')
    add_full_table(doc, ['カテゴリ', '要件', '目標値', '計測方法'], [
        ['セキュリティ', 'APIキーフロント露出', '0件', 'bandit自動スキャン（CI/CD）'],
        ['セキュリティ', 'PII暗号化（企業名・担当者名・連絡先）', 'AES-256 + RLS', 'bandit + 定期監査'],
        ['セキュリティ', 'ISMS/ISO27001準拠', '情報セキュリティ管理体制整備', '内部監査 / 外部審査'],
        ['性能', 'APIレスポンス速度', '0.7秒以下', 'Cloud Trace / Chrome DevTools'],
        ['性能', 'マッチング処理時間（バッチ）', '1,000件 / 5分以内', 'Cloud Functions実行ログ'],
        ['コスト', '月額インフラコスト（MVP）', '¥5,000以下', 'GCP請求 / 月次'],
        ['可用性', 'Cloud Run SLA', '99.95%以上', 'GCP SLAに依存'],
        ['オフライン', 'IndexedDB niceeze_cache_v142 動作保証',
         'ネットワーク切断時も在庫一覧閲覧可能', '統合テスト'],
        ['拡張性', '同時接続ユーザー数', '1,000社（Phase 1）→ 10,000社（Phase 2）', 'Cloud Run自動スケール'],
    ])

    # 5. データモデル
    add_heading(doc, '5. データモデル（Firestore コレクション設計）')
    add_full_table(doc, ['エンティティ', 'コレクション名', '主要フィールド', '保存先'], [
        ['SurplusItem', 'surplus_items',
         'item_id (str) / seller_id (str) / category (str) / quantity (int) / '
         'purchase_price (int) / desired_price (int) / status (str) / created_at (ts)',
         'Firestore / IndexedDB v142'],
        ['BuyerProfile', 'buyer_profiles',
         'buyer_id (str) / company_name_encrypted (str) / preferred_categories (list) / '
         'region (str) / budget_max (int) / min_quantity (int) / updated_at (ts)',
         'Firestore / IndexedDB v142'],
        ['MatchRecord', 'match_records',
         'match_id (str) / item_id (str) / buyer_id (str) / seller_id (str) / '
         'match_score (float) / status (str) / negotiation_deadline (ts) / agreed_price (int)',
         'Firestore'],
        ['TransactionLog', 'transaction_logs',
         'tx_id (str) / match_id (str) / contract_url (str) / sha256_hash (str) / '
         'completed_at (ts) / amount (int)',
         'Firestore / BigQuery（月次アーカイブ）'],
        ['AuditLog', 'audit_logs',
         'id (uuid) / event_type (str) / actor_id (str) / payload (map) / '
         'sha256_hash (str) / created_at (ts)',
         'Firestore（append-only RLS）'],
    ])

    # 6. API仕様
    add_heading(doc, '6. API仕様（Cloud Run エンドポイント一覧）')
    add_full_table(doc, ['エンドポイント', 'メソッド', '概要', '認証', '実装Gate'], [
        ['/api/surplus/items', 'POST', '余剰在庫登録', 'Bearer JWT', 'G1'],
        ['/api/surplus/items', 'GET', '余剰在庫一覧取得（フィルタ付き）', 'Bearer JWT', 'G1'],
        ['/api/surplus/items/{item_id}', 'PATCH', '在庫ステータス更新', 'Bearer JWT', 'G1'],
        ['/api/matching/run', 'POST', 'マッチング実行（非同期）', 'Bearer JWT', 'G1'],
        ['/api/matching/{match_id}', 'GET', 'マッチング結果取得', 'Bearer JWT', 'G1'],
        ['/api/negotiation/{match_id}/price', 'POST', '価格交渉提案送信', 'Bearer JWT', 'G2'],
        ['/api/transaction/confirm', 'POST', '取引確定・契約書生成', 'Bearer JWT', 'G2'],
        ['/api/reports/summary', 'GET', '転換実績サマリーレポート', 'Bearer JWT', 'G3'],
        ['/api/admin/health', 'GET', 'ヘルスチェック', 'Service Account', 'G0'],
    ])

    # 7. セキュリティ要件
    add_heading(doc, '7. セキュリティ要件（bandit / PII取扱方針）')
    add_body(doc,
        'banditセキュリティスキャンは全Gateで必須。bandit検知件数は0件が必須条件であり、'
        '1件でも検知された場合はCI/CDビルドを即時停止する。'
        'PII（個人情報）は企業担当者名・メールアドレス・電話番号が対象であり、'
        'AES-256暗号化（pgcrypto）+ PostgreSQL Row Level Securityで保護する。'
    )
    add_full_table(doc, ['要件', '実装方式', '対象', '確認方法'], [
        ['banditスキャン 0件必須', 'GitHub Actions CI（bandit -r . -ll）', '全Pythonソース', 'CI自動実行 / Gate必須条件'],
        ['PII暗号化', 'AES-256（pgcrypto encrypt_pii/decrypt_pii）', '担当者名・連絡先', 'bandit + 統合テスト'],
        ['行レベルセキュリティ', 'PostgreSQL RLS（company_id一致のみ参照可）', '全テーブル', 'SQLテスト'],
        ['APIキー管理', 'GCP Secret Manager（フロント露出0件）', '全外部API認証情報', 'bandit自動スキャン'],
        ['契約書改ざん防止', 'SHA-256ハッシュ署名 + タイムスタンプ', '生成契約書PDF/Word', '監査ログ照合'],
        ['ISMS準拠', 'ISO27001情報セキュリティ管理体制', 'システム全体', '内部監査 / 外部審査（G4）'],
        ['監査ログ', 'append-only（DELETE禁止RLS）+ SHA-256署名', '全システム操作', '監査レポート'],
    ])

    # 8. 技術スタック
    add_heading(doc, '8. 技術スタック定義')
    add_full_table(doc, ['レイヤー', '技術', 'バージョン/設定', '用途'], [
        ['フロントエンド', 'HTML5 / Vanilla JS', 'ES2022+', '売り手ダッシュボード / 買い手ポータル'],
        ['フロントエンド', 'IndexedDB', 'v142（niceeze_cache_v142）', 'オフラインキャッシュ'],
        ['モバイル', 'PWA', 'Service Worker / Web App Manifest', 'モバイル対応'],
        ['バックエンド', 'Python / FastAPI', '3.11 / 0.110+', 'Cloud Run APIサーバー'],
        ['バックエンド', 'python-docx', '1.x', '契約書Word生成'],
        ['バックエンド', 'Cloud Functions', 'Python 3.11', 'マッチング非同期処理'],
        ['DB', 'Firestore', 'ネイティブモード', 'リアルタイム在庫・マッチングデータ'],
        ['DB', 'Cloud SQL PostgreSQL', 'AES-256 + RLS', 'PII保護・監査ログ'],
        ['ストレージ', 'Cloud Storage', 'Standard', '契約書PDF/Word保管'],
        ['分析', 'BigQuery', 'Standard SQL', '月次アーカイブ・レポート'],
        ['セキュリティ', 'GCP Secret Manager', 'v1', 'APIキー管理'],
        ['CI/CD', 'GitHub Actions', 'bandit-check.yml', 'セキュリティスキャン必須'],
    ])

    doc.save(OUTPUT)
    print(f"✅ {OUTPUT} 生成完了")


if __name__ == '__main__':
    build()
