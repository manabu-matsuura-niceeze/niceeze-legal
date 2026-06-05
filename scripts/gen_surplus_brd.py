"""SURPLUS_SHIFT_BRD_v1.0.docx 生成スクリプト"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = "docs/SURPLUS_SHIFT/SURPLUS_SHIFT_BRD_v1.0.docx"

DARK_NAVY = RGBColor(0x1a, 0x3a, 0x5c)
ORANGE    = RGBColor(0xf5, 0xa6, 0x23)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF0, 0xF4, 0xF8)


def set_margins(doc, top=20, bottom=20, left=20, right=20):
    sec = doc.sections[0]
    sec.top_margin    = Mm(top)
    sec.bottom_margin = Mm(bottom)
    sec.left_margin   = Mm(left)
    sec.right_margin  = Mm(right)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12) if level == 1 else Pt(10.5)
    run.font.color.rgb = color or DARK_NAVY
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_table_width(table, width_dxa=9026):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def add_full_table(doc, headers, rows, header_bg='1a3a5c', header_fg=WHITE):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_width(table, 9026)
    table.style = 'Table Grid'

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = header_fg

    for ri, row_data in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F0F4F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)

    return table


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            'SURPLUS_SHIFT 余剰在庫転換システム  |  ビジネス要件定義書 v1.0  |  © 2026 株式会社NiceEze  Confidential'
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)


def build():
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc = Document()
    set_margins(doc)
    add_footer(doc)

    # ─── 表紙ブロック ───────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SURPLUS_SHIFT — ビジネス要件定義書 (BRD) v1.0')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('余剰在庫転換システム（SURPLUS_SHIFT）　v1.0').font.size = Pt(12)

    doc.add_paragraph()

    # ─── 1. 文書管理表 ────────────────────────────────────────
    add_heading(doc, '1. 文書管理表')
    add_full_table(doc,
        ['項目', '内容'],
        [
            ['文書名', 'SURPLUS_SHIFT ビジネス要件定義書'],
            ['バージョン', 'v1.0'],
            ['作成日', '2026-06-05'],
            ['最終更新日', '2026-06-05'],
            ['作成者', 'NiceEze 自律COO'],
            ['承認者', '代表取締役CEO 松浦 学'],
            ['ステータス', '承認済（Gate 1スプリント適用中）'],
            ['関連文書', 'SURPLUS_SHIFT_SRS_v1.0.docx / PLAN-20260605-001'],
        ]
    )

    # ─── 2. ビジネス背景 ──────────────────────────────────────
    add_heading(doc, '2. ビジネス背景（余剰在庫問題）')
    add_body(doc,
        'EC仕入れ業者・製造業者において、需要予測の誤差や季節変動により余剰在庫が慢性的に発生している。'
        '余剰在庫は保管コスト・廃棄コストの増大を招き、キャッシュフロー悪化・機会損失の主要因となっている。'
        '特に中小規模のEC事業者は適切な買い手とのマッチングチャネルを持たず、'
        '不良在庫を廃棄または大幅値引きで処分せざるを得ない状況が続いている。'
        'SURPLUS_SHIFTは余剰在庫の自動検知・買い手マッチング・転換実行を統合し、'
        '在庫転換率70%以上・平均マッチング時間48時間以内を目標とする。'
    )

    add_full_table(doc,
        ['課題', '現状値', 'SURPLUS_SHIFT目標値'],
        [
            ['余剰在庫の平均滞留期間', '90日以上', '30日以内（自動マッチング）'],
            ['手動マッチング工数', '1件あたり8〜24時間', '完全自動（0工数）'],
            ['在庫廃棄率', '15〜30%', '5%以下'],
            ['在庫転換率', '40〜50%', '70%以上（KPI目標）'],
        ]
    )

    # ─── 3. 主要機能一覧 ──────────────────────────────────────
    add_heading(doc, '3. 主要機能一覧')
    add_full_table(doc,
        ['ID', '機能名', '概要', '優先度'],
        [
            ['SURPLUS-001', '売り手ダッシュボード',
             '余剰在庫の一覧表示・登録・ステータス管理。在庫数量・仕入価格・希望転換価格を入力',
             'Must'],
            ['SURPLUS-002', '買い手検索マッチング',
             '登録済み買い手プロファイルとのAI自動マッチング。カテゴリ・地域・予算・数量条件でスコアリング',
             'Must'],
            ['SURPLUS-003', '価格交渉支援',
             '売り手・買い手間の価格交渉を自動化。市場価格参照・推奨価格提示・交渉履歴記録',
             'Should'],
            ['SURPLUS-004', '転換実行・契約書生成',
             'マッチング成立後の取引確定・契約書自動生成（PDF/Word）・決済フロー連携',
             'Must'],
            ['SURPLUS-005', '履歴・分析レポート',
             '転換実績・収益改善額・在庫削減率のレポート生成。BigQueryによる月次分析',
             'Should'],
        ]
    )

    # ─── 4. ステークホルダー ──────────────────────────────────
    add_heading(doc, '4. ステークホルダー定義')
    add_full_table(doc,
        ['ステークホルダー', '役割', '主要タッチポイント', '優先度'],
        [
            ['売り手企業', '余剰在庫の登録・転換依頼者', 'SURPLUS-001 売り手ダッシュボード（Web）', 'Must'],
            ['買い手企業', '余剰在庫の購入希望者', 'SURPLUS-002 買い手検索・マッチング画面', 'Must'],
            ['NiceEzeオペレーター', 'システム監視・例外処理・KPI確認', 'COO報告パネル / Cloud Console', 'Must'],
            ['GCP管理者', 'インフラ運用・コスト管理', 'Cloud Console / Secret Manager', 'Should'],
        ]
    )

    # ─── 5. 成功指標（KPI） ───────────────────────────────────
    add_heading(doc, '5. 成功指標（KPI）')
    add_full_table(doc,
        ['KPI', '目標値', '計測方法', 'Gate'],
        [
            ['余剰在庫転換率', '≥70%', '転換完了件数 / 登録総件数', 'G2'],
            ['平均マッチング時間', '≤48時間', 'マッチング成立TSから計算', 'G1'],
            ['ユーザー満足度', '≥4.0/5.0', 'アプリ内評価フォーム', 'G2'],
            ['APIキーフロント露出件数', '0件', 'bandit自動スキャン', '全Gate'],
            ['月額インフラコスト（MVP）', '¥5,000以下', 'GCP請求 / 月次', 'G1'],
            ['在庫廃棄率', '≤5%', '廃棄処理件数 / 登録総件数', 'G3'],
            ['契約書生成エラー率', '0%', '統合テスト / 本番監査ログ', 'G2'],
        ]
    )

    # ─── 6. Gate制・FinOps ────────────────────────────────────
    add_heading(doc, '6. Gate制（G0〜G4）/ FinOps 月額上限管理')
    add_full_table(doc,
        ['Gate', '完了条件', 'FinOps上限', '備考'],
        [
            ['G0', 'GCP環境構築・Secret Manager設定完了', '¥0（無料枠）', 'インフラ基盤のみ'],
            ['G1', 'SURPLUS-001/002 基本動作確認・bandit 0件', '¥5,000/月（MVP上限）',
             'Cloud Run + Firestore無料枠活用'],
            ['G2', 'SURPLUS-003/004 動作確認・KPI初期計測開始', '¥10,000/月', '価格交渉・契約書生成'],
            ['G3', 'SURPLUS-005 レポート機能・BigQuery連携完了', '¥15,000/月', '分析基盤'],
            ['G4', '3社以上の実証実験完了・転換率70%達成', '¥30,000/月', 'スケールアップ準備'],
        ]
    )

    add_body(doc,
        'FinOps原則: Cloud Run・Firestore無料枠を最大活用。'
        'MVP（G1完了）時点での月額インフラコスト上限は¥5,000。'
        'GCP予算アラートを¥4,000（80%）および¥5,000（100%）に設定し、超過時は自動通知。'
    )

    # ─── 7. 制約条件・前提条件 ────────────────────────────────
    add_heading(doc, '7. 制約条件・前提条件')
    add_full_table(doc,
        ['区分', '内容'],
        [
            ['技術制約', 'APIキーはフロントエンドに一切露出しない（DevSecOps）'],
            ['技術制約', 'GCPサーバーレス構成（Cloud Run / Firestore / Cloud Functions）固定'],
            ['技術制約', 'オフライン動作: IndexedDB niceeze_cache_v142 使用'],
            ['法的制約', '個人情報保護法: PII（企業名・担当者名・連絡先）はAES-256暗号化＋RLS'],
            ['法的制約', '商取引法: 電子契約書はタイムスタンプ付与・改ざん防止対応必須'],
            ['業務制約', '売り手・買い手双方の企業認証（法人番号照合）を必須とする'],
            ['前提条件', 'GCP Secret Managerにシークレット設定済み（Gate 0）'],
            ['前提条件', '法人番号API（国税庁）への接続環境が整備済み'],
        ]
    )

    # ─── 不滅憲章チェックリスト ───────────────────────────────
    add_heading(doc, '8. 不滅憲章との整合性確認チェックリスト')
    add_full_table(doc,
        ['原則', '確認項目', '判定'],
        [
            ['隠蔽禁止', 'APIキーをフロントエンドに含めていないか', '✅ 確認済'],
            ['隠蔽禁止', 'エラー情報を握り潰していないか（監査ログに記録）', '✅ 確認済'],
            ['不明点隔離', '未確定仕様を【松浦CEO要件定義待ち】として明記しているか', '✅ 確認済'],
            ['顧客起点', 'KPIが売り手・買い手企業の体験向上に直結しているか', '✅ 確認済'],
            ['Gate制', 'Gate 0〜G1の未承認実装を含んでいないか', '✅ 確認済'],
            ['FinOps', 'MVP月額¥5,000上限を遵守しているか', '✅ 確認済'],
        ]
    )

    doc.save(OUTPUT)
    print(f"✅ {OUTPUT} 生成完了")


if __name__ == '__main__':
    build()
