"""SBDS_SRS_v1.0.docx 生成スクリプト"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "docs/SBDS/SBDS_SRS_v1.0.docx"
DARK_NAVY = RGBColor(0x1a, 0x3a, 0x5c)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)


def set_margins(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Mm(20)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_table_width(table, width_dxa=9026):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12) if level == 1 else Pt(10.5)
    run.font.color.rgb = DARK_NAVY


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)


def add_full_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_width(table, 9026)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hdr.cells[i], '1a3a5c')
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F0F4F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            set_cell_bg(tr.cells[ci], bg)
            run = tr.cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('館内配送システム（SBDS）  |  システム要件定義書 v1.0  |  © 2026 株式会社NiceEze  Confidential')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)


def build():
    doc = Document()
    set_margins(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('システム要件定義書（SRS）')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('館内配送システム（SBDS）　v1.0').font.size = Pt(12)

    # 1. 文書管理表
    add_heading(doc, '1. 文書管理表')
    add_full_table(doc, ['項目', '内容'], [
        ['文書名', 'SBDS システム要件定義書'],
        ['バージョン', 'v1.0'],
        ['作成日', '2026-06-05'],
        ['作成者', '自律COO（Claude Code）'],
        ['承認者', '代表取締役CEO 松浦 学'],
        ['根拠文書', 'Ver.14.2要件定義書 / BRD v1.0 / LAYOUT_MASTER.md v1.0'],
    ])

    # 2. システム全体像
    add_heading(doc, '2. システム全体像・対象領域')
    add_body(doc,
        'SBDSは集合住宅館内配送を完全デジタル化するシステムである。'
        '管理画面（TMS-SET-001）で建物マスタを設定し、配送員スマホアプリ（TMS-DRV-001）で'
        '最適ルート配送を実現する。LINE LIFF/PWAによるモバイルファーストUI、'
        'IndexedDB（niceeze_cache_v142）によるオフライン動作、'
        'GCP Cloud Run + Firestoreによるサーバーレス構成で運用する。'
    )
    add_full_table(doc, ['コンポーネント', '役割', '技術'], [
        ['TMS-SET-001', '建物マスタ・フロアグリッドエディタ', 'HTML/JS / IndexedDB v142 / Firestore'],
        ['TMS-DRV-001', '配送員スマホ・ルーティング画面', 'PWA / LINE LIFF / WebXR / IndexedDB v142'],
        ['Cloud Run', 'APIサーバー / LINEプロキシ', 'Python / FastAPI'],
        ['Firestore', 'リアルタイムDB（配送データ）', 'GCP Firestore'],
        ['Memorystore Redis', 'LINE Webhookデデュープ', 'Redis Streams Consumer Group'],
        ['BigQuery', '月次配送データアーカイブ', 'bigquery_pipeline.py'],
        ['Secret Manager', 'APIキー管理', 'GCP Secret Manager'],
    ])

    # 3. 5つの絶対条件
    add_heading(doc, '3. 5つの絶対条件（Ver.14.2不変制約）')
    add_full_table(doc, ['#', '絶対条件', '具体的要件', '違反時の対応'], [
        ['1', 'APIキーフロント露出禁止',
         'ANTHROPIC_API_KEY / LINE_CHANNEL_SECRET等は全てCloud Run経由またはSecret Manager経由のみ',
         'bandit検知後即時ビルド停止'],
        ['2', 'IndexedDB v142固定',
         'DB名: niceeze_cache_v142 / IDB_VERSION: 142（v140からの移行完了）',
         'バージョン不一致はデータ不採用'],
        ['3', 'Jaro-Winkler閾値 D_jw≥0.85',
         '宛名名寄せ閾値は0.85固定。変更はCEO承認必須',
         'テスト失敗でデプロイ停止'],
        ['4', '労働法ロック 4時間',
         'LABOR_LAW_BREAK_MINUTES=240固定。STATUS_LOCKED_BY_LABOR_LAWに強制遷移',
         'UI上で入力操作を完全無効化'],
        ['5', 'LAYOUT_MASTER.md準拠',
         '画面レイアウト変更はCEO承認→LAYOUT_MASTER更新→実装の順序厳守',
         'layout-guard CIがビルドを拒否'],
    ])

    # 4. 機能要件一覧
    add_heading(doc, '4. 機能要件一覧')
    add_full_table(doc,
        ['ID', '機能名', '詳細', '優先度', '画面ID'],
        [
            ['SRS-S-001', '建物マスタ登録', '棟数(1-20)・階数(1-100)・居住者用EV(0-20基)・業務用EV(最低4基)の入力', 'Must', 'TMS-SET-001'],
            ['SRS-S-002', 'フロアグリッド編集', '棟名/部屋番号/専有面積(㎡)/家賃(円)/EV出口距離(m)/階のスプレッドシート型入力', 'Must', 'TMS-SET-001'],
            ['SRS-S-003', 'DXFインポート', 'dxf-parser.jsによるブラウザサイドDXF解析・RoomRecord自動生成', 'Must', 'TMS-SET-001'],
            ['SRS-S-004', 'AR計測', 'WebXR Hit Test APIによるEV出口距離実測（±50cm許容）', 'Should', 'TMS-SET-001'],
            ['SRS-S-005', 'IndexedDB v142保存', 'フロアグリッドデータをniceeze_cache_v142に保存・読み込み', 'Must', '両画面'],
            ['SRS-S-006', '最適ルーティングソート', 'フロア昇順→EV出口距離昇順→クレーム要注意最後尾のソート', 'Must', 'TMS-DRV-001'],
            ['SRS-S-007', 'Jaro-Winkler名寄せ', '宛名表記ゆれD_jw≥0.85で同一人物判定', 'Must', 'TMS-DRV-001'],
            ['SRS-S-008', '冷凍/冷蔵警告', '冷凍・冷蔵フラグ時に赤字大文字「ロッカー格納禁止：手渡し必須」を表示', 'Must', 'TMS-DRV-001'],
            ['SRS-S-009', '1分前PULL通知', 'estimated_arrival_ts - now ≤ 60秒でLINE PULL通知発火', 'Must', 'TMS-DRV-001'],
            ['SRS-S-010', '労働法ロック', '4時間(240分)連続でSTATUS_LOCKED_BY_LABOR_LAWに遷移・全入力無効', 'Must', 'TMS-DRV-001'],
            ['SRS-S-011', 'ETA残り時間表示', '推定到着残り時間をリアルタイム（1秒更新）でfont-mono表示', 'Must', 'TMS-DRV-001'],
            ['SRS-S-012', 'Firestore永続化', 'BuildingMasterをFirestoreのbuilding_mastersコレクションに保存', 'Must', 'バックエンド'],
            ['SRS-S-013', 'LINE Webhookデデュープ', 'Redis Streams Consumer GroupによるLINE Webhook重複排除', 'Must', 'バックエンド'],
            ['SRS-S-014', 'BigQuery月次アーカイブ', '配送完了データを月次でBigQueryへアーカイブ', 'Should', 'バックエンド'],
        ]
    )

    # 5. 非機能要件
    add_heading(doc, '5. 非機能要件')
    add_full_table(doc, ['カテゴリ', '要件', '目標値', '計測方法'], [
        ['性能', 'IndexedDB応答速度', '0.7秒以下', 'Chrome DevTools / Lighthouse'],
        ['性能', 'Cloud Run API応答速度', '500ms以下', 'GCP Cloud Trace'],
        ['コスト', '1個口インフラコスト（3万世帯）', '¥0.50以下（実績¥0.10〜0.14）', 'GCP請求/月次配送件数'],
        ['可用性', 'Cloud Run SLA', '99.95%以上', 'GCP SLAに依存'],
        ['セキュリティ', 'PII暗号化', 'AES-256（pgcrypto）+ Row Level Security', 'bandit + 定期監査'],
        ['セキュリティ', 'LINE Webhook署名検証', 'HMAC-SHA256（X-Line-Signature）', '統合テスト'],
        ['セキュリティ', 'APIキーフロント露出', '0件', 'bandit自動スキャン（layout-guard CI）'],
        ['拡張性', '最大世帯数', '3万世帯（Phase 1）→ 10万世帯（Phase 2）', 'Cloud Run自動スケール'],
        ['オフライン', 'IndexedDB動作保証', 'ネットワーク切断時も配送リスト閲覧・完了登録可能', '統合テスト'],
    ])

    # 6. 技術スタック
    add_heading(doc, '6. 技術スタック定義')
    add_full_table(doc, ['レイヤー', '技術', 'バージョン/設定', '用途'], [
        ['フロントエンド', 'HTML5 / Vanilla JS', 'ES2022+', 'TMS-SET-001 / TMS-DRV-001'],
        ['フロントエンド', 'IndexedDB', 'v142（niceeze_cache_v142）', 'オフラインキャッシュ'],
        ['フロントエンド', 'WebXR Device API', 'Level 1 / Hit Test', 'AR計測（判断③適用後）'],
        ['フロントエンド', 'dxf-parser.js', 'MIT License', 'DXFファイル解析'],
        ['モバイル', 'LINE LIFF', 'v2.x', '配送員スマホUI'],
        ['モバイル', 'PWA', 'Service Worker / Web App Manifest', 'LIFF非対応時フォールバック'],
        ['バックエンド', 'Python / FastAPI', '3.11 / 0.110+', 'Cloud Run APIサーバー'],
        ['DB', 'Firestore', 'ネイティブモード', 'リアルタイム配送データ'],
        ['DB', 'Cloud SQL PostgreSQL', 'AES-256 + RLS + パーティション', 'スケール後移行'],
        ['キャッシュ', 'Memorystore Redis', '7.x', 'LINE Webhookデデュープ'],
        ['分析', 'BigQuery', 'Standard SQL', '月次アーカイブ'],
        ['セキュリティ', 'GCP Secret Manager', 'v1', 'APIキー管理'],
        ['CI/CD', 'GitHub Actions', 'layout-guard.yml', 'レイアウトガバナンス'],
    ])

    # 7. 外部API
    add_heading(doc, '7. 外部API・連携システム一覧')
    add_full_table(doc, ['API/システム', '用途', '認証方式', '実装Gate'], [
        ['LINE Messaging API', '1分前PULL通知 / Webhook受信', 'Channel Secret + Access Token', 'G1'],
        ['LINE LIFF', '配送員スマホUI', 'LIFF ID', 'G1'],
        ['WebXR Hit Test API', 'AR距離計測', 'ブラウザネイティブ（認証不要）', 'G1（判断③後）'],
        ['GCP Firestore', '配送データ永続化', 'Service Account IAM', 'G0'],
        ['GCP Secret Manager', 'APIキー管理', 'Service Account IAM', 'G0'],
        ['GCP BigQuery', '月次アーカイブ', 'Service Account IAM', 'G1'],
        ['Claude API（将来）', 'ルーティング最適化AI', 'Cloud Run Proxy経由', 'G4'],
    ])

    # 8. データモデル概要
    add_heading(doc, '8. データモデル概要')
    add_full_table(doc, ['エンティティ', 'キーフィールド', '主要フィールド', '保存先'], [
        ['BuildingMaster', 'property_id (str)', 'spec, rooms[], created_at, updated_at', 'Firestore / IndexedDB v142'],
        ['RoomRecord', '(building_name, room_number)', 'area_sqm, rent_jpy, ev_exit_distance_m, floor', 'Firestore / IndexedDB v142'],
        ['DeliveryPackage', 'package_id (str)', 'recipient_name, floor, is_frozen, is_refrigerated, is_complaint_risk, status, estimated_arrival_ts', 'Firestore / IndexedDB v142'],
        ['WorkSession', 'staff_id (str)', 'start_ts, status（ACTIVE/STATUS_LOCKED_BY_LABOR_LAW）', 'IndexedDB v142'],
        ['AuditLog', 'id (uuid)', 'event_type, payload, sha256_hash, created_at', 'Cloud SQL（append-only RLS）'],
    ])

    # 9. セキュリティ要件
    add_heading(doc, '9. セキュリティ要件')
    add_full_table(doc, ['要件', '実装方式', '対象', '確認方法'], [
        ['PII暗号化', 'AES-256（pgcrypto encrypt_pii/decrypt_pii）', '氏名・住所', 'bandit + 統合テスト'],
        ['行レベルセキュリティ', 'PostgreSQL RLS（staff_id一致のみ参照可）', '全テーブル', 'SQLテスト'],
        ['LINE署名検証', 'HMAC-SHA256（X-Line-Signature）', 'LINE Webhook受信', '統合テスト'],
        ['APIキー管理', 'GCP Secret Manager（フロント露出0件）', '全外部API', 'bandit自動スキャン'],
        ['Redis重複排除', 'Consumer Group XACK/XCLAIM（処理タイムアウト30秒）', 'LINE Webhook', '負荷テスト'],
        ['監査ログ', 'append-only（DELETE禁止RLS）+ SHA-256署名', '全システム操作', '監査レポート'],
    ])

    # 未決事項
    add_heading(doc, '【松浦CEO要件定義待ち】未決事項')
    add_full_table(doc, ['判断ID', '内容', '影響範囲'], [
        ['判断③', 'AR計測対象デバイス（iOS/Android同時対応 or Android優先）',
         'WebXR実装工数・G1完了時期'],
    ])

    doc.save(OUTPUT)
    print(f"✅ {OUTPUT} 生成完了")


if __name__ == '__main__':
    build()
